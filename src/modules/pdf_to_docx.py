import pdfplumber
from docx import Document
from docx.shared import Pt
import os


class PdfToDocxConverter:
    """Convert PDF files to DOCX format."""

    def __init__(self, input_path: str):
        self.input_path = input_path

    def convert(self, output_path: str, start_page: int = 0, end_page: int = None, multi_page_table: bool = True) -> dict:
        """Convert PDF to DOCX.
        
        Args:
            output_path: Path to save DOCX file
            start_page: Starting page (0-indexed)
            end_page: Ending page (None for all pages)
            multi_page_table: Parse tables spanning multiple pages
        
        Returns:
            dict with conversion stats
        """
        doc = Document()
        
        with pdfplumber.open(self.input_path) as pdf:
            pages = pdf.pages[start_page:end_page]
            
            for page in pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        if line.strip():
                            p = doc.add_paragraph(line)
                            p.style.font.size = Pt(11)
                
                if multi_page_table:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            doc_table = doc.add_table(rows=len(table), cols=len(table[0]))
                            for i, row in enumerate(table):
                                for j, cell in enumerate(row):
                                    doc_table.rows[i].cells[j].text = str(cell) if cell else ""
        
        doc.save(output_path)
        
        return {
            "input_file": self.input_path,
            "output_file": output_path,
            "output_size": os.path.getsize(output_path)
        }

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass
